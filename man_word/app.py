from docx import Document

nome_arquivo = './teste.docx'  # Nome do arquivo Word
texto_antigo = 'nome_'  # Texto que deve ser substituído
texto_novo = 'Raphael Fernanades Franca com '  # Novo texto para substituir

documento = Document(nome_arquivo) #abrindo o arquivo

# Substitui o texto antigo pelo novo em todo o parágrafo

'''
for p in range(10):
    paragrafo = documento.paragraphs[p]  # Assumindo que o texto está no primeiro parágrafo

    paragrafo.text = paragrafo.text.replace(texto_antigo, texto_novo)

    # Salva o documento modificado
    
'''

for paragrafo in documento.paragraphs:
    for linha in paragrafo.runs:
        if texto_antigo in linha.text:
            # Substituir o texto na linha
            linha.text = linha.text.replace(texto_antigo, texto_novo)
            # Marcar como encontrado (opcional)
            linha.font.bold = True
            documento.save('nome_arquivo.docx')
